#!/usr/bin/env python3
"""
Build a single Word .docx distribution pack for "The Library Is Now an App"
— all channel posts (LinkedIn / Facebook / X / Reddit) in one paste-ready doc
with real clickable hyperlinks and image checklists per channel.

Output: docs/exports/blog-the-library-is-now-an-app-DISTRIBUTION-PACK.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SUBSTACK_URL = "https://lindsayhiebert.substack.com/p/the-library-is-now-an-app"
PRODUCT_URL = "https://aistockassist.com"
LINKEDIN_PROFILE = "https://www.linkedin.com/in/lindsayhiebert/"
SUBSTACK_HOME = "https://lindsayhiebert.substack.com/"

IMAGE_DIR = "docs/images/library-is-now-an-app/"


def add_hyperlink(paragraph, url, text):
    """Insert a real clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    return h


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p


def divider(doc):
    p = doc.add_paragraph("━" * 40)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    return p


def callout_label(doc, text):
    """Yellow-ish bold label that marks 'PASTE THIS' blocks."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x6B, 0x00)
    run.font.size = Pt(11)
    return p


def paste_block(doc, lines):
    """A clearly demarcated paste block using Normal style — zero indent,
    default font, no quote styling. Bracketed by visible START/END COPY
    markers so the user can highlight between them. This guarantees clean
    plain-text paste into LinkedIn / FB / X without leading whitespace."""
    # ── Top marker (NOT part of the copy region) ──
    top = doc.add_paragraph("✂  ─────────  START COPY HERE  ─────────  ✂")
    top.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    top.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    top.runs[0].font.size = Pt(9)
    top.runs[0].bold = True

    # ── Body: Normal style paragraphs, no indent, default font ──
    for line in lines:
        p = doc.add_paragraph()  # Normal — zero indent, default Calibri
        if line.strip():
            p.add_run(line)
        # else: empty paragraph for vertical spacing

    # ── Bottom marker (NOT part of the copy region) ──
    bot = doc.add_paragraph("✂  ─────────   END COPY HERE   ─────────  ✂")
    bot.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    bot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    bot.runs[0].font.size = Pt(9)
    bot.runs[0].bold = True
    doc.add_paragraph()  # trailing spacer


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──────────────────────────────────────────────────────────
    title = doc.add_heading("The Library Is Now an App", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = sub.add_run("Multi-Channel Distribution Pack")
    sr.italic = True
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dr = date_p.add_run("Published Thu May 21, 2026  ·  Amplification window: Thu May 21 → Mon May 25")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    divider(doc)

    # ── How to copy ────────────────────────────────────────────────────
    h2(doc, "How to copy without breaking the paste")
    para(doc, "Every post in this pack is wrapped between two visible markers:")
    p = doc.add_paragraph()
    p.add_run("    ✂  ───  START COPY HERE  ───  ✂").italic = True
    p.add_run("\n    (your post text)\n    ").italic = True
    p.add_run("✂  ───   END COPY HERE   ───  ✂").italic = True
    bullets(doc, [
        "Highlight ONLY the lines BETWEEN the two ✂ markers — NOT the markers themselves.",
        "All paste blocks use Normal paragraph style (zero left indent, default font) — so the text copies cleanly with no leading whitespace.",
        "Em-dashes (—), arrows (↑ ↓ →), icons (📚 🚀 ❌ ✅ ⚠️ 👉 👇 📖 ⚡), and # hashtags all render fine on LinkedIn / Facebook / X — they are standard Unicode.",
        "If you paste into LinkedIn and see odd formatting, use the LinkedIn editor's 'Clear formatting' (Tx) button — but with these blocks you shouldn't need to.",
    ])
    doc.add_paragraph()
    para(doc, "Style note — icons by channel:", bold=True)
    bullets(doc, [
        "LinkedIn / Facebook / X — written in the icon-rich, line-broken, ❌/🚀-contrasted style that drives engagement on visual feeds.",
        "Reddit — INTENTIONALLY plain prose, no decorative emoji. r/investing, r/SecurityAnalysis, r/ValueInvesting, r/Bogleheads treat heavy-emoji posts as marketing/spam and downvote them.",
    ])

    divider(doc)

    # ── Quick reference ────────────────────────────────────────────────
    h2(doc, "Quick reference")

    p = doc.add_paragraph()
    p.add_run("Substack article (canonical):  ").bold = True
    add_hyperlink(p, SUBSTACK_URL, SUBSTACK_URL)

    p = doc.add_paragraph()
    p.add_run("Product / try-the-tool:  ").bold = True
    add_hyperlink(p, PRODUCT_URL, PRODUCT_URL)

    p = doc.add_paragraph()
    p.add_run("LinkedIn profile:  ").bold = True
    add_hyperlink(p, LINKEDIN_PROFILE, LINKEDIN_PROFILE)

    p = doc.add_paragraph()
    p.add_run("Image folder:  ").bold = True
    p.add_run(IMAGE_DIR).font.name = "Consolas"

    divider(doc)

    # ── Schedule ──────────────────────────────────────────────────────
    h2(doc, "Posting schedule")
    schedule = [
        ("✅ Thu May 21 ~8am",  "Substack",            "DONE — live"),
        ("⏳ Thu May 21 +30m",  "LinkedIn carousel",   "Caption + first comment"),
        ("⏳ Thu May 21 ~6pm",  "Facebook",            "Multi-image + caption"),
        ("⏳ Thu May 21 ~9pm",  "X / Twitter",         "7-tweet thread"),
        ("⏳ Fri May 22 9am",   "Reddit r/investing",  "PRIMARY drop — block 2-6 hrs for replies"),
        ("⏳ Fri May 22",       "Medium",              "Passive cross-post"),
        ("⏳ Sat May 23 9am",   "Reddit r/SecurityAnalysis", "Academic angle"),
        ("⏳ Sun May 24 9am",   "Reddit r/ValueInvesting",   "'600 books later' angle"),
        ("⏳ Mon May 25 9am",   "Reddit r/Bogleheads",       "Framework discipline angle"),
        ("⏳ Mon May 25 6pm",   "Reddit r/financialindependence", "FIRE / long-term-thinking angle"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "When"
    hdr[1].text = "Channel"
    hdr[2].text = "Format"
    for when, ch, fmt in schedule:
        row = table.add_row().cells
        row[0].text = when
        row[1].text = ch
        row[2].text = fmt
    doc.add_paragraph()
    divider(doc)

    # ════════════════════════════════════════════════════════════════════
    # LINKEDIN CAROUSEL
    # ════════════════════════════════════════════════════════════════════
    h1(doc, "1️⃣  LinkedIn — Carousel (PRIMARY)")
    para(doc, "Why carousel: ~3-5× the reach of single-image. Substack URL goes in the FIRST COMMENT — not in the post body — to preserve LinkedIn algorithm reach.", italic=True)

    h3(doc, "Carousel images — upload in this order (1080×1080 PDF export from Canva)")
    carousel_images = [
        ("SLIDE 1 — Hook",          "LIBRARY_HERO.png + text overlay: \"I read 600+ investing books.\" / \"Then I built the tool I wished existed.\""),
        ("SLIDE 2 — The Problem",   "LIBRARY_INARTICLE_books-and-app.png + overlay: \"The wisdom of Buffett, Graham, Munger exists in books.\" / \"The problem was never access. The problem was translation.\""),
        ("SLIDE 3 — The Engine",    "LIBRARY_INARTICLE_jargon-decryption.png + overlay: \"AI Stock Assist turns Wall Street jargon into plain English.\""),
        ("SLIDE 4 — The Proof",     "LIBRARY_INARTICLE_business-vs-story-APA.png + overlay: \"This week I tested 6 tickers.\" / \"The framework disagreed with Wall Street twice.\""),
        ("SLIDE 5 — Differentiator","LIBRARY_INARTICLE_greats-boardroom-CNQ.png + overlay: \"Cash flow. Valuation. Quality. Risk. Discipline.\""),
        ("SLIDE 6 — CTA",           "LIBRARY_CAROUSEL_slide-6-cta.png (NEEDS GENERATION — prompt in blog .md L213). Headline: \"3 free analyses.\" Footer: → aistockassist.com"),
    ]
    for slide, spec in carousel_images:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(slide + " — ").bold = True
        p.add_run(spec)

    callout_label(doc, "📋 PASTE THIS — Caption (below the carousel)")
    paste_block(doc, [
        "I have analyzed 600+ investing books since 2008. 📚",
        "",
        "One question kept coming back: why doesn't every ordinary investor have this wisdom at their fingertips?",
        "",
        "The problem was never access. The books exist.",
        "The problem was translation.",
        "",
        "So I built AI Stock Assist.",
        "",
        "❌ Not a chatbot.",
        "❌ Not another stock screener.",
        "",
        "🚀 A working analytical engine that applies the disciplined frameworks of Buffett, Graham, Munger, Lynch, Bogle, and Housel to real-time market data — and gives you back a verdict in plain English.",
        "",
        "This week I tested six tickers.",
        "The framework disagreed with Wall Street twice:",
        "",
        "One had remarkably strong cash flow.",
        "",
        "The other was burning cash with negative free cash flow.",
        "",
        "Wall Street prices the story.",
        "The framework tests the business.",
        "",
        "👇 How to check the results:",
        "",
        "👉 Swipe through the slides ↑ for the visual story.",
        "👉 Read the full breakdown on Substack — link is in the first comment below! ↓",
        "",
        "#Investing #ValueInvesting #BenjaminGraham #WarrenBuffett #FintechAI #StockMarket #DividendInvesting",
    ])

    callout_label(doc, "📋 PASTE THIS — First comment (5 min after the carousel goes live)")
    paste_block(doc, [
        f"Full longread on Substack: {SUBSTACK_URL}",
        "",
        f"Try the tool — 3 free analyses, no credit card: {PRODUCT_URL}",
    ])

    divider(doc)

    # ════════════════════════════════════════════════════════════════════
    # LINKEDIN SINGLE-IMAGE FALLBACK
    # ════════════════════════════════════════════════════════════════════
    h1(doc, "1️⃣b  LinkedIn — Single-image FALLBACK (skip if you posted the carousel)")
    h3(doc, "Image to attach")
    bullets(doc, ["LIBRARY_HERO.png  (or LIBRARY_PROMO_linkedin.png as an alternate composition)"])

    callout_label(doc, "📋 PASTE THIS — Post body (~290 words)")
    paste_block(doc, [
        "I have analyzed 600+ investing books since 2008. 📚",
        "",
        "Have you ever dreamed of investing with the discipline of Buffett, Graham, Munger, Lynch, Dalio, Bogle, and Housel?",
        "",
        "I have.",
        "",
        "So I built the library.",
        "",
        "One question would not leave me alone:",
        "👉 Why doesn't every ordinary investor have this wisdom at their fingertips?",
        "",
        "The problem was never access. The books exist.",
        "The problem was translation.",
        "",
        "That is why I built AI Stock Assist.",
        "",
        "❌ Not a chatbot.",
        "❌ Not another stock screener.",
        "",
        "🚀 A working analytical engine that applies timeless investment frameworks to real-time data — and turns finance jargon into a verdict you can understand in seconds.",
        "",
        "This week I analyzed six tickers: T · BMY · CNQ · HBM · APA · EGO.",
        "",
        "In sixty seconds, AI Stock Assist disagreed with Wall Street twice.",
        "",
        "Not contrarian for the sake of it. The framework simply saw what consensus often misses:",
        "",
        "✅ One stock — Wall Street said HOLD. AI Stock Assist saw strong free cash flow, low P/FCF, and the kind of cash conversion discipline Graham and Buffett would care about.",
        "",
        "⚠️ Another — Wall Street said HOLD with major implied upside. AI Stock Assist saw negative free cash flow, weak cash generation, and a story the numbers did not support.",
        "",
        "Wall Street prices the story.",
        "AI Stock Assist tests the business.",
        "",
        "Cash flow. Valuation. Quality. Risk. Discipline.",
        "",
        "Built with the experience of the investing greats. Amplified by frontier AI on the cutting edge of better investment intelligence.",
        "",
        f"📖 Read the full piece: {SUBSTACK_URL}",
        "",
        f"🚀 3 free analyses, no credit card → {PRODUCT_URL}",
        "",
        "P.S. If you've read even one of the books I mentioned and felt the same itch — the gap between what's written and what's applied — I'd love your reply. What was the book that changed how you invest?",
        "",
        "#Investing #ValueInvesting #BenjaminGraham #WarrenBuffett #PersonalFinance #FintechAI #ArtificialIntelligence #StockMarket #AIforGood #DividendInvesting #FinancialEducation",
    ])

    divider(doc)

    # ════════════════════════════════════════════════════════════════════
    # FACEBOOK
    # ════════════════════════════════════════════════════════════════════
    h1(doc, "2️⃣  Facebook — Multi-image post (~6pm ET)")
    para(doc, "Facebook does NOT penalize external links in the body. Upload 5–6 images as a single multi-photo post — auto-becomes a swipeable gallery on mobile.", italic=True)

    h3(doc, "Images to upload (in this order)")
    bullets(doc, [
        "LIBRARY_HERO.png",
        "LIBRARY_INARTICLE_books-and-app.png",
        "LIBRARY_INARTICLE_jargon-decryption.png",
        "LIBRARY_INARTICLE_business-vs-story-APA.png",
        "LIBRARY_INARTICLE_greats-boardroom-CNQ.png",
    ])

    callout_label(doc, "📋 PASTE THIS — Caption (~210 words)")
    paste_block(doc, [
        "I want to share something I've been working on for years. 📚",
        "",
        "Since 2008 I've read or listened to over 600 books on investing — Buffett, Graham, Munger, Lynch, Dalio, Bogle, Housel, all of them.",
        "",
        "One question kept nagging at me:",
        "👉 Why doesn't every ordinary investor have this wisdom at their fingertips?",
        "",
        "The books exist. The problem was never access — it was translation.",
        "",
        "So I built it. AI Stock Assist.",
        "",
        "❌ Not a chatbot.",
        "❌ Not another screener.",
        "",
        "🚀 A working analytical engine that applies the same disciplined frameworks the greats used — to real-time data — and turns finance jargon into a verdict you can understand in seconds.",
        "",
        "This week I tested six tickers: T, BMY, CNQ, HBM, APA, EGO.",
        "",
        "Sixty seconds later, AI Stock Assist disagreed with Wall Street twice.",
        "",
        "✅ Strong cash flow on one.",
        "⚠️ Weak cash generation on another.",
        "",
        "Not contrarian. Just disciplined.",
        "",
        "Wall Street prices the story.",
        "AI Stock Assist tests the business.",
        "",
        f"📖 Read the full piece: {SUBSTACK_URL}",
        "",
        f"🚀 3 free analyses, no credit card → {PRODUCT_URL}",
        "",
        "#Investing #PersonalFinance #StockMarket",
    ])

    divider(doc)

    # ════════════════════════════════════════════════════════════════════
    # X / TWITTER
    # ════════════════════════════════════════════════════════════════════
    h1(doc, "3️⃣  X / Twitter — 7-tweet thread (~9pm ET)")
    para(doc, "Post as one connected thread — reply each tweet to the previous one.", italic=True)

    tweets = [
        ("1/", [
            "Have you ever dreamed of investing with the discipline of Buffett, Graham, Munger, Lynch, Dalio, Bogle, and Housel? 📚",
            "",
            "I have.",
            "",
            "So I built the library.",
        ]),
        ("2/", [
            "Since 2008 I've read or listened to 600+ books on investing, finance, economics, business, psychology, risk, and human behavior. 📚",
            "",
            "One question kept coming back:",
            "",
            "👉 Why doesn't every ordinary investor have this wisdom at their fingertips?",
        ]),
        ("3/", [
            "The problem was never access. The books exist.",
            "",
            "The problem was translation.",
            "",
            "So I built AI Stock Assist.",
            "",
            "❌ Not a chatbot.",
            "❌ Not a screener.",
            "",
            "🚀 A working analytical engine that applies timeless investment frameworks to real-time data.",
        ]),
        ("4/", [
            "This week I ran six tickers through it:",
            "",
            "T · BMY · CNQ · HBM · APA · EGO",
            "",
            "⚡ In 60 seconds, AI Stock Assist disagreed with Wall Street twice.",
            "",
            "Not contrarian. Just disciplined.",
        ]),
        ("5/", [
            "✅ On one — Wall Street said HOLD.",
            "The framework saw strong free cash flow, low P/FCF, and the cash conversion discipline Graham and Buffett would care about.",
            "",
            "⚠️ On another — Wall Street said HOLD with upside.",
            "The framework saw negative FCF and a story the numbers didn't support.",
        ]),
        ("6/", [
            "Wall Street prices the story.",
            "AI Stock Assist tests the business.",
            "",
            "Cash flow. Valuation. Quality. Risk. Discipline.",
            "",
            "Built with the experience of the investing greats. Amplified by frontier AI.",
            "",
            "A framework that does not flinch.",
        ]),
        ("7/", [
            "📚 The library is now an app.",
            "",
            "🚀 3 free analyses, no credit card.",
            "",
            f"→ {PRODUCT_URL}",
            "",
            f"📖 Full piece: {SUBSTACK_URL}",
            "",
            "#Investing #AI #ValueInvesting",
        ]),
    ]
    for num, lines in tweets:
        callout_label(doc, f"📋 PASTE THIS — Tweet {num}")
        paste_block(doc, lines)

    divider(doc)

    # ════════════════════════════════════════════════════════════════════
    # REDDIT
    # ════════════════════════════════════════════════════════════════════
    h1(doc, "4️⃣  Reddit — 5 subs over 4 days")

    para(doc,
         "Rules of engagement (every drop):",
         bold=True)
    bullets(doc, [
        "NO link in the post body. Link in profile only. Parenthetical mention OK at the end.",
        "Lead with the INSIGHT, not the product.",
        "Block 2–6 hrs after posting for replies. Reddit's algorithm rewards early engagement velocity.",
        "Reply to every comment within 2–6 hrs for the first 24 hrs.",
        "Stagger 24–48 hrs between subs — same-day reposts get spam-flagged.",
        "Don't argue downvotes — let it breathe.",
    ])
    doc.add_paragraph()

    # ── r/investing
    h2(doc, "r/investing — Fri May 22, 9am ET  (PRIMARY DROP)")
    para(doc, "Broadest sub — lead with concrete numbers.", italic=True)

    callout_label(doc, "📋 PASTE THIS — Title")
    paste_block(doc, ["I've read 600+ investing books. Here's what kept bugging me — and what I built."])

    callout_label(doc, "📋 PASTE THIS — Body")
    paste_block(doc, [
        "Long-time lurker. Have read or listened to over 600 books on investing, finance, economics, business, psychology, risk, and human behavior since 2008. Buffett, Graham, Munger, Lynch, Dalio, Bogle, Housel — all the canon, plus a lot of the second tier.",
        "",
        "One question kept coming back to me: why doesn't every ordinary investor have this wisdom available in a usable form? The books exist — the problem was never access. The problem was translation. You either spend a decade reading them yourself, or you pay a wealth manager who may or may not actually apply the frameworks.",
        "",
        "So I built AI Stock Assist. It's not a chatbot. It's not another screener. It's an analytical engine that applies the same disciplined frameworks the greats used — to real-time market data — and produces a verdict in 60 seconds with the reasoning grounded in the actual numbers.",
        "",
        "I tested it on six tickers this week: T, BMY, CNQ, HBM, APA, EGO. The framework disagreed with the Wall Street consensus on two of them. Not because I told it to be contrarian — because Wall Street tends to price the story (analyst targets, momentum, sentiment) while a Graham/Buffett/Munger framework tests the business (FCF, P/FCF, OCF/NI, payout sustainability, balance sheet quality).",
        "",
        "I'm not pitching anything in this post — just sharing the philosophical itch that drove the build. The link is in my profile if you want to see it. 3 free analyses, no signup required. Mostly I'd love feedback on whether the verdict logic resonates with how you all think about quality vs price.",
    ])

    divider(doc)

    # ── r/SecurityAnalysis
    h2(doc, "r/SecurityAnalysis — Sat May 23, 9am ET")
    para(doc, "Academic / Graham-literate. Lead with the translation-not-access framing.", italic=True)

    callout_label(doc, "📋 PASTE THIS — Title")
    paste_block(doc, ["The problem with the Graham/Dodd canon was never access — it was translation. (What I built to fix that.)"])

    callout_label(doc, "📋 PASTE THIS — Body (academic tone-shift of the r/investing post)")
    paste_block(doc, [
        "Spent the last 15 years working through the canon — Security Analysis, The Intelligent Investor, Common Stocks and Uncommon Profits, Margin of Safety, Poor Charlie's Almanack — and a few hundred others. The frameworks are clear. The discipline is clear. What's missing for the ordinary investor isn't access to the texts — it's translation from text to repeatable application on real-time data.",
        "",
        "That's what I built. AI Stock Assist applies the disciplined frameworks of Graham, Dodd, Buffett, Munger, Lynch, Bogle, and Housel to live market data — and renders a verdict in 60 seconds with the reasoning grounded in the actual numbers (FCF, P/FCF, OCF/NI, payout ratio, dividend coverage, balance sheet quality, Graham number, margin of safety).",
        "",
        "I ran six tickers this week as a sanity check: T, BMY, CNQ, HBM, APA, EGO. The framework disagreed with Wall Street consensus on two of them. Wall Street is largely pricing the story (analyst targets, momentum, sentiment). A disciplined value/quality framework tests the business.",
        "",
        "I'd genuinely value feedback from this sub on the verdict logic. The link's in my profile. 3 free analyses, no signup. What I'm most curious about: does the synthesis hold up under the same scrutiny you'd apply to a Graham screen run by hand?",
    ])

    divider(doc)

    # ── r/ValueInvesting
    h2(doc, "r/ValueInvesting — Sun May 24, 9am ET")
    para(doc, "Practical value-investor tone. Lead with '600 books later, here's the gap I built for.'", italic=True)

    callout_label(doc, "📋 PASTE THIS — Title")
    paste_block(doc, ["600 books later, here's the gap I built for."])

    callout_label(doc, "📋 PASTE THIS — Body")
    paste_block(doc, [
        "Long-time value-investing nerd. Read or listened to 600+ books on investing since 2008 — Graham, Buffett, Munger, Lynch, Klarman, Greenblatt, the whole canon plus the second tier. One thing kept bugging me: the frameworks work. The discipline works. But the gap between reading the books and actually applying the screens to a live ticker, in real time, in plain English — that gap was huge for the ordinary investor.",
        "",
        "So I built AI Stock Assist. It applies the same disciplined screens the greats used — to real-time data — and produces a verdict in 60 seconds. FCF, P/FCF, OCF/NI, Graham number, margin of safety, payout sustainability, balance sheet quality. The whole stack.",
        "",
        "Tested six tickers this week: T, BMY, CNQ, HBM, APA, EGO. The framework disagreed with Wall Street twice. One had strong free cash flow + low P/FCF + healthy cash conversion despite a Wall Street HOLD. Another had negative FCF and weak cash generation despite a Wall Street HOLD-with-upside.",
        "",
        "Not contrarian for fun. Just disciplined. Link in profile, 3 free analyses, no signup. Mostly looking for feedback from people who actually read the canon — does the verdict logic feel like it's applying the same discipline you'd apply by hand?",
    ])

    divider(doc)

    # ── r/Bogleheads
    h2(doc, "r/Bogleheads — Mon May 25, 9am ET")
    para(doc, "⚠️ Skeptical of stock-picking. Lead with framework discipline, NOT picks. Frame as a discipline-check tool, not a stock recommender.", italic=True)

    callout_label(doc, "📋 PASTE THIS — Title")
    paste_block(doc, ["A discipline check, not a stock-picker: applying the index-investing canon (Bogle, Malkiel, Ellis) to individual-stock temptations."])

    callout_label(doc, "📋 PASTE THIS — Body")
    paste_block(doc, [
        "Boglehead by conviction — 90% of my portfolio is in index funds and always will be. But a lot of us still have the temptation, or the family pressure, or the 'one stock from work' problem. The question isn't whether to pick individual stocks (the answer is mostly no). The question is: if you ARE going to look at one, how do you apply the same discipline the canon teaches?",
        "",
        "I read 600+ books on investing since 2008. Bogle, Malkiel, Ellis, Bernstein — the index-investing case is settled. But I also read Graham, Buffett, Munger, Klarman. What I kept noticing: the discipline overlap is huge. Low fees, long horizons, margin of safety, cash flow quality, ignore the story, focus on the business.",
        "",
        "So I built AI Stock Assist as a discipline check. Not a recommender. You hand it a ticker; it applies a Graham/Buffett/Munger framework to live data; it tells you in plain English whether the BUSINESS holds up — FCF, OCF/NI, payout sustainability, balance sheet quality, Graham number. 60 seconds.",
        "",
        "Tested six tickers this week. Framework disagreed with Wall Street consensus on two. Not because it's contrarian — because Wall Street prices the story; the framework tests the business.",
        "",
        "Link in profile. 3 free analyses, no signup. Genuinely curious whether the Boglehead crowd sees this as a useful discipline check for the stock-picking temptations we all occasionally face, or whether the framing itself bothers you. Either feedback is valuable.",
    ])

    divider(doc)

    # ── r/financialindependence
    h2(doc, "r/financialindependence — Mon May 25, 6pm ET")
    para(doc, "FIRE crowd. Long-horizon thinking + book canon + the discipline angle.", italic=True)

    callout_label(doc, "📋 PASTE THIS — Title")
    paste_block(doc, ["After 600 books on investing, the FIRE lesson that mattered most was 'translation' — not 'access.'"])

    callout_label(doc, "📋 PASTE THIS — Body")
    paste_block(doc, [
        "Pursuing FIRE since 2008. Read or listened to over 600 books along the way — Bogle, Graham, Buffett, Munger, Lynch, Housel, Collins, Sethi, all of them. The one lesson that became impossible to ignore: the books are not the bottleneck. The frameworks have been free and public for decades. The bottleneck for the ordinary investor is translation — turning what's written into something you can actually apply to a real ticker in real time.",
        "",
        "So I built AI Stock Assist. It applies the disciplined frameworks the greats used — to real-time market data — and gives you a verdict in 60 seconds. Cash flow quality, valuation, balance sheet health, payout sustainability, margin of safety. In plain English.",
        "",
        "Why it matters for FIRE specifically: the long horizon makes mistakes expensive. A wrong dividend-stock pick at year 5 of a 25-year glide path compounds for 20 more years. A disciplined framework check at the buy decision is the cheapest insurance you can get.",
        "",
        "Ran six tickers this week: T, BMY, CNQ, HBM, APA, EGO. Framework disagreed with Wall Street twice — including one stock the consensus rates HOLD with implied upside, but where the cash flow tells a different story.",
        "",
        "Link in profile. 3 free analyses, no signup. Mostly looking for FIRE-community feedback: does the long-horizon-discipline framing resonate, or does it feel like one more shiny tool that distracts from the boring index-fund truth?",
    ])

    divider(doc)

    # ════════════════════════════════════════════════════════════════════
    # MEDIUM + WEBSITE
    # ════════════════════════════════════════════════════════════════════
    h1(doc, "5️⃣  Medium — Passive cross-post")
    para(doc, "Open Medium → New Story → paste the entire Substack body (URL-pasted; Medium will auto-embed). Title: 'The Library Is Now an App.' Add tags: investing, value-investing, ai, personal-finance, financial-education. Set canonical URL → Substack URL.", italic=False)

    p = doc.add_paragraph()
    p.add_run("Canonical URL to set in Medium settings:  ").bold = True
    add_hyperlink(p, SUBSTACK_URL, SUBSTACK_URL)

    divider(doc)

    h1(doc, "6️⃣  Website cross-link (already drafted, not yet deployed)")
    bullets(doc, [
        "src/components/LearnPage.tsx — Library added to featuredArticles[0] with tag 'Latest' (May 21, 2026)",
        "Status: code change saved locally; NOT yet committed/pushed",
        "Deploy: git commit + git push (Vercel auto-deploys on push to main)",
    ])

    divider(doc)

    # ── Footer
    f = doc.add_paragraph()
    f.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = f.add_run("— End of distribution pack —")
    fr.italic = True
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    out = Path("docs/exports/blog-the-library-is-now-an-app-DISTRIBUTION-PACK.docx").resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
