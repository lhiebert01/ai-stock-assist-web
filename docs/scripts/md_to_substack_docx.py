#!/usr/bin/env python3
"""
Convert a blog episode markdown file into a Substack-friendly .docx.

- Strips the "# Channel Variants" section and everything after (Substack longread only)
- Renders headings as Word headings
- Preserves bold, italic, and links inline
- Image markers (`> **[IMAGE: ...]**`) become italic figure placeholders
- Captions (`> *Caption: ...*`) become italic caption lines
- Code-fenced blocks (the THE MATH callout) become monospace blocks
- Horizontal rules become a centered divider

Usage:
    python3 md_to_substack_docx.py INPUT.md [OUTPUT.docx]

If OUTPUT is omitted, writes to the same basename in docs/exports/.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


INLINE_RE = re.compile(
    r"\*\*(?P<bold>[^*]+)\*\*"          # **bold**
    r"|\*(?P<italic>[^*]+)\*"           # *italic*
    r"|\[(?P<link>[^\]]+)\]\((?P<url>[^)]+)\)"  # [text](url)
    r"|(?P<plain>(?:[^*\[\\]|\\.)+)"    # plain text
)


def add_inline(paragraph, text):
    """Append runs to `paragraph` with inline markdown handled."""
    for m in INLINE_RE.finditer(text):
        if m.group("bold"):
            run = paragraph.add_run(m.group("bold"))
            run.bold = True
        elif m.group("italic"):
            run = paragraph.add_run(m.group("italic"))
            run.italic = True
        elif m.group("link"):
            run = paragraph.add_run(m.group("link"))
            run.font.color.rgb = RGBColor(0x06, 0x4E, 0xCC)
            run.underline = True
        elif m.group("plain"):
            paragraph.add_run(m.group("plain"))


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text()

    # Drop everything from "# Channel Variants" onward
    if "# Channel Variants" in text:
        text = text.split("# Channel Variants")[0]

    doc = Document()

    # Set base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    lines = text.split("\n")
    in_code = False
    code_buffer = []
    table_rows = []
    in_table = False

    def flush_table():
        """
        Render a markdown table as bulleted paragraphs instead of a Word
        table — Substack and Medium both strip Word tables on paste.

        Pattern per data row:
          • **First cell (bold)** · Header2: value2 · Header3: value3 · *Last cell (italic)*
        """
        nonlocal table_rows, in_table
        if not table_rows:
            return

        # Single-row "table" → render as a plain paragraph
        if len(table_rows) < 2:
            p = doc.add_paragraph()
            for i, cell in enumerate(table_rows[0]):
                if i > 0:
                    p.add_run(" · ")
                add_inline(p, cell)
            doc.add_paragraph()
            table_rows = []
            in_table = False
            return

        headers = table_rows[0]
        data_rows = table_rows[1:]

        for row in data_rows:
            p = doc.add_paragraph(style="List Bullet")
            n = len(row)

            # First cell — bold (the row name / company)
            if n >= 1:
                first_run = p.add_run(row[0])
                first_run.bold = True

            # Middle cells — "Header: value" joined by " · "
            for i in range(1, n - 1):
                p.add_run(" · ")
                label = headers[i].strip() if i < len(headers) else ""
                if label:
                    p.add_run(f"{label}: ")
                add_inline(p, row[i])

            # Last cell — italic (often the verdict / summary)
            if n >= 2:
                p.add_run(" · ")
                last_label = headers[-1].strip() if len(headers) >= n else ""
                # Skip the label if it's a generic "Verdict" / "Result" — value speaks for itself
                if last_label and last_label.lower() not in {"verdict", "result", "outcome", "rating"}:
                    p.add_run(f"{last_label}: ")
                last_run = p.add_run(row[-1])
                last_run.italic = True

        doc.add_paragraph()  # spacing after the "table" block
        table_rows = []
        in_table = False

    for line in lines:
        rstripped = line.rstrip()

        # Code block start/end (renders as monospace block — used for THE MATH callout)
        if rstripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                code_buffer = []
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        # Markdown table detection: a line starting with `|`
        if rstripped.startswith("|"):
            cells = [c.strip() for c in rstripped.strip("|").split("|")]
            # Skip the separator row (---|---|---)
            if all(re.match(r"^:?-+:?$", c) for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()

        # Horizontal rule
        if rstripped == "---":
            p = doc.add_paragraph("⸻")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue

        # Headings
        if rstripped.startswith("# "):
            doc.add_heading(rstripped[2:], level=1)
            continue
        if rstripped.startswith("## "):
            doc.add_heading(rstripped[3:], level=2)
            continue
        if rstripped.startswith("### "):
            doc.add_heading(rstripped[4:], level=3)
            continue
        if rstripped.startswith("#### "):
            doc.add_heading(rstripped[5:], level=4)
            continue

        # Image marker:  > **[IMAGE: ...]**
        m_img = re.match(r"^> \*\*\[IMAGE:\s*(.+?)\]\*\*", rstripped)
        if m_img:
            p = doc.add_paragraph()
            run = p.add_run(f"[ Image: {m_img.group(1)} ]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Caption line:  > *Caption: ...*
        m_cap = re.match(r"^> \*Caption:\s*(.+?)\*\s*$", rstripped)
        if m_cap:
            p = doc.add_paragraph()
            run = p.add_run("Caption: " + m_cap.group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Generic blockquote
        if rstripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, rstripped[2:])
            continue

        # Bullet list
        if rstripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, rstripped[2:])
            continue

        # Numbered list
        m_num = re.match(r"^\d+\.\s+(.*)", rstripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m_num.group(1))
            continue

        # Empty line
        if not rstripped:
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline(p, rstripped)

    if in_table:
        flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 md_to_substack_docx.py INPUT.md [OUTPUT.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1]).resolve()
    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2]).resolve()
    else:
        # Default: docs/exports/<slug>.docx
        repo_docs = md_path.parent
        docx_path = repo_docs / "exports" / (md_path.stem + ".docx")

    convert(md_path, docx_path)


if __name__ == "__main__":
    main()
