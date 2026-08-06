#!/usr/bin/env python3
"""Build the .docx twin of a distribution-pack .md (WO-ASA-BLOG-NVO-001 T6).

Same pattern as the MacroLens pack builder: the .md is the single source of
truth and the .docx is ALWAYS regenerated from it — never hand-edited — so the
two can't drift. Handles the markdown subset the packs use: #/## headings,
**bold** / *italic* / `code` inline, _italic paragraph_ notes, the ▼▲ copy
markers (rendered as shaded marker lines), and bare URLs.

Usage:
    python3 scripts/build_pack_docx.py docs/social/2026-08-06_every-signal-was-green_pack.md
    # writes the .docx next to the .md
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent

ACCENT = RGBColor(0x08, 0x91, 0xB2)   # cyan-700 — readable on white
MUTED = RGBColor(0x64, 0x74, 0x8B)
MARKER = RGBColor(0x0E, 0x74, 0x90)

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(par, text: str) -> None:
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = par.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
        else:
            par.add_run(part)


def build(md_path: Path) -> Path:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("**Hero image (ATTACH THIS FILE):**"):
            # Render the line, then EMBED the hero right below it so the doc is
            # fully self-contained — the poster saves/drag-drops it from here.
            p = doc.add_paragraph()
            add_inline(p, stripped)
            m = re.search(r"public/\S+\.(?:jpg|png)", stripped)
            hero = REPO / m.group(0) if m else None
            if hero and hero.exists():
                pic = doc.add_paragraph()
                pic.add_run().add_picture(str(hero), width=Inches(6.0))
                cap = doc.add_paragraph()
                r = cap.add_run("^ THE HERO — right-click → Save as Picture (or drag out) to attach on LinkedIn. Also in Downloads as asa-every-signal-was-green-1200x630.jpg.")
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = MUTED
            else:
                warn = doc.add_paragraph()
                r = warn.add_run("[hero image not found in repo — attach from the live URL instead]")
                r.italic = True
                r.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("▼") or stripped.startswith("▲"):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.bold = True
            r.font.color.rgb = MARKER
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif stripped.startswith("_") and stripped.endswith("_"):
            p = doc.add_paragraph()
            r = p.add_run(stripped[1:-1])
            r.italic = True
            r.font.color.rgb = MUTED
            r.font.size = Pt(10)
        elif re.fullmatch(r"https?://\S+", stripped):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.color.rgb = ACCENT
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    if len(sys.argv) != 2:
        sys.exit("usage: build_pack_docx.py <pack.md>")
    src = Path(sys.argv[1])
    print(f"wrote {build(src)}")
